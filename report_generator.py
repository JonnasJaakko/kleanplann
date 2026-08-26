"""Production-safe DOCX report generator for KleanPlann."""
from __future__ import annotations

import io
from pathlib import Path

import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
import matplotlib.patches as mpatches
import numpy as np
from PIL import Image

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from cost_calculator import calculate_cost
from schedule_validator import validate_schedule


class ReportGenerationError(RuntimeError):
    pass


def _hex_color(rgb): return "#{:02X}{:02X}{:02X}".format(*[int(x) for x in rgb[:3]])


def _set_cell_shading(cell, color_hex):
    tc_pr=cell._tc.get_or_add_tcPr(); old=tc_pr.find(qn("w:shd"))
    if old is not None: tc_pr.remove(old)
    el=OxmlElement("w:shd"); el.set(qn("w:fill"),color_hex.replace("#","")); el.set(qn("w:val"),"clear"); tc_pr.append(el)


def _set_repeat_header(row):
    trpr=row._tr.get_or_add_trPr(); el=OxmlElement("w:tblHeader"); el.set(qn("w:val"),"true"); trpr.append(el)


def _set_cant_split(row):
    trpr=row._tr.get_or_add_trPr(); el=OxmlElement("w:cantSplit"); trpr.append(el)


def _set_cell_margins(cell, top=60, start=70, bottom=60, end=70):
    tc_pr=cell._tc.get_or_add_tcPr(); mar=tc_pr.first_child_found_in("w:tcMar")
    if mar is None: mar=OxmlElement("w:tcMar"); tc_pr.append(mar)
    for side,val in (("top",top),("start",start),("bottom",bottom),("end",end)):
        node=mar.find(qn(f"w:{side}"))
        if node is None: node=OxmlElement(f"w:{side}"); mar.append(node)
        node.set(qn("w:w"),str(val)); node.set(qn("w:type"),"dxa")


def _style_header(row):
    _set_repeat_header(row)
    for cell in row.cells:
        _set_cell_shading(cell,"1F4E78"); _set_cell_margins(cell)
        for p in cell.paragraphs:
            p.alignment=WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs: r.bold=True; r.font.color.rgb=RGBColor(255,255,255); r.font.size=Pt(8)


def _protect_table(table):
    for row in table.rows[1:]:
        _set_cant_split(row)
        for cell in row.cells: _set_cell_margins(cell)


def _duration(minutes):
    minutes=int(round(minutes)); return f"{minutes//60} ч {minutes%60:02d} мин ({minutes/60:.2f} ч)"


def _find_room(project, room_id, floor_index=None):
    if floor_index is not None and 0<=floor_index<len(project.floors):
        r=next((x for x in project.floors[floor_index].rooms if x.id==room_id),None)
        if r: return r
    return next((x for f in project.floors for x in f.rooms if x.id==room_id),None)


def _zone_for(project, floor_index, room_id):
    for z in project.zones:
        if getattr(z,"floor_index",0)==floor_index and room_id in z.room_ids: return z
    for z in project.zones:
        if room_id in z.room_ids: return z
    return None


def _floor_area(floor):
    return float(getattr(floor,"total_area_m2",0) or sum(r.area_m2 for r in floor.rooms))


def _resolve_image_path(project,floor):
    raw=str(getattr(floor,"image_path","") or "").strip()
    if not raw: return None
    p=Path(raw); candidates=[p] if p.is_absolute() else [Path(getattr(project,"_project_dir","") or ".")/p,Path.cwd()/p,Path(__file__).resolve().parent/p]
    for c in candidates:
        try:
            if c.exists() and c.is_file(): return c
        except OSError: pass
    return None


def _draw_floor_plan(project,floor,employee_index=None,dpi=180):
    image_path=_resolve_image_path(project,floor); arr=None
    if image_path:
        try: arr=np.array(Image.open(image_path).convert("RGB"))
        except Exception as exc: raise ReportGenerationError(f"Не удалось прочитать исходный план '{image_path}': {exc}") from exc
    elif getattr(floor,"image_path",""):
        raise ReportGenerationError(f"Не найден исходный план этажа '{floor.name}'.\nПуть: {floor.image_path}\nЭкспорт остановлен, чтобы не выпустить неполный документ.")
    if arr is not None:
        sw,sh=arr.shape[1],arr.shape[0]
    else:
        pts=[p for r in floor.rooms for p in r.points]
        xs=[p[0] for p in pts] if pts else [0,1000]; ys=[p[1] for p in pts] if pts else [0,700]
        sw=max(xs)-min(xs) or 1; sh=max(ys)-min(ys) or 1
    ratio=max(.45,min(2.2,float(sw)/float(sh))); fw=8.1; fh=max(4.2,min(7.0,fw/ratio))
    fig,ax=plt.subplots(figsize=(fw,fh),dpi=dpi)
    if arr is not None:
        ax.imshow(arr,extent=[0,sw,sh,0],origin="upper",alpha=.43,interpolation="nearest"); ax.set_xlim(0,sw); ax.set_ylim(sh,0)
    else:
        pts=[p for r in floor.rooms for p in r.points]
        if pts:
            xs=[p[0] for p in pts]; ys=[p[1] for p in pts]; mx=max(10,(max(xs)-min(xs))*.04); my=max(10,(max(ys)-min(ys))*.04); ax.set_xlim(min(xs)-mx,max(xs)+mx); ax.set_ylim(max(ys)+my,min(ys)-my)
    for room in floor.rooms:
        z=_zone_for(project,floor.index,room.id); rgb=tuple(z.color[:3]) if z else tuple(room.color[:3]); active=employee_index is None or (z and z.employee_index==employee_index); alpha=.76 if active else .10
        ax.add_patch(mpatches.Polygon(room.points,closed=True,facecolor=[x/255 for x in rgb],alpha=alpha,edgecolor="black",linewidth=.7))
        if room.points:
            cx=sum(x for x,_ in room.points)/len(room.points); cy=sum(y for _,y in room.points)/len(room.points)
            if employee_index is None and z: ax.text(cx,cy,str(z.employee_index+1),ha="center",va="center",fontsize=8,fontweight="bold",color="white",bbox=dict(facecolor=_hex_color(rgb),edgecolor="black",pad=1.0))
            if employee_index is None or (z and z.employee_index==employee_index): ax.text(cx,cy,f"№{room.id+1}",ha="center",va="center",fontsize=5,bbox=dict(facecolor="white",alpha=.75,pad=.6))
    ax.set_aspect("equal"); ax.axis("off")
    buf=io.BytesIO(); fig.savefig(buf,format="png",bbox_inches="tight",pad_inches=.04); plt.close(fig); buf.seek(0); return buf


def _fit_picture(stream,max_width,max_height):
    pos=stream.tell(); stream.seek(0); with_image=Image.open(stream); w,h=with_image.size; stream.seek(pos)
    scale=min(max_width*180/w,max_height*180/h)
    return {"width":Inches(w*scale/180),"height":Inches(h*scale/180)}


def _set_margins(section,top=.45,bottom=.45,left=.55,right=.55):
    section.top_margin=Inches(top); section.bottom_margin=Inches(bottom); section.left_margin=Inches(left); section.right_margin=Inches(right)


def _new_landscape(doc,project):
    sec=doc.add_section(WD_SECTION.NEW_PAGE); sec.orientation=WD_ORIENT.LANDSCAPE; sec.page_width=Inches(11); sec.page_height=Inches(8.5); _set_margins(sec,.45,.45,.55,.55); _add_footer(sec,project); return sec


def _add_footer(section,project):
    # Sections created with python-docx inherit the previous footer by default.
    # Unlink it before writing, otherwise every new section appends its title to
    # the same shared paragraph (a very visible DOCX production defect).
    section.footer.is_linked_to_previous = False
    footer = section.footer
    p = footer.paragraphs[0]
    p.clear()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(f"KleanPlann • {project.name}")
    r.font.size = Pt(8)
    r.font.color.rgb = RGBColor(120,120,120)


def _set_styles(doc):
    doc.styles["Normal"].font.name="Aptos"; doc.styles["Normal"].font.size=Pt(9)
    for name,size,color in (("Title",22,RGBColor(31,78,120)),("Heading 1",15,RGBColor(31,78,120)),("Heading 2",12,RGBColor(47,84,150)),("Heading 3",10,RGBColor(68,68,68))):
        s=doc.styles[name]; s.font.name="Aptos"; s.font.size=Pt(size); s.font.color.rgb=color


def _employee_stats(project,emp):
    tasks=sorted([t for t in project.cleaning_tasks if t.employee==emp],key=lambda t:t.start_dt); minutes=sum((t.end_dt-t.start_dt).total_seconds()/60 for t in tasks); overtime=sum((t.end_dt-t.start_dt).total_seconds()/60 for t in tasks if getattr(t,"is_overtime",False)); transit=sum(float(getattr(t,"transit_before_minutes",0)) for t in tasks); rooms={(t.floor_index,t.room_id) for t in tasks}; area=sum((_find_room(project,rid,fi).area_m2 if _find_room(project,rid,fi) else 0) for fi,rid in rooms); return tasks,minutes,overtime,transit,rooms,area


def _add_schedule_table(doc,tasks,project):
    floorcol=len(project.floors)>1; headers=(['Этаж'] if floorcol else [])+['№','Комната','Тип','Площадь, м²','Начало','Конец','Мин.']
    table=doc.add_table(rows=1,cols=len(headers)); table.style="Table Grid"
    for i,h in enumerate(headers): table.rows[0].cells[i].text=h
    _style_header(table.rows[0])
    for t in sorted(tasks,key=lambda x:x.start_dt):
        room=_find_room(project,t.room_id,t.floor_index); row=table.add_row(); _set_cant_split(row); vals=([str(t.floor_index+1)] if floorcol else [])+[str(t.room_id+1),room.name if room else "—",room.room_type if room else "—",f"{room.area_m2:.1f}" if room else "—",f"{t.start_dt:%H:%M}",f"{t.end_dt:%H:%M}",str(int(round((t.end_dt-t.start_dt).total_seconds()/60)))]
        for i,v in enumerate(vals): row.cells[i].text=v
        if getattr(t,"is_overtime",False):
            for c in row.cells: _set_cell_shading(c,"F4CCCC")
    _protect_table(table); return table


def _add_summary_table(doc,cost):
    rows=[("Время уборки",_duration(cost.get("cleaning_minutes",0))),("Переходы",_duration(cost.get("transit_minutes",0))),("Общее рабочее время",_duration(cost.get("total_minutes",0))),("Переработка",_duration(cost.get("overtime_hours",0)*60) if cost.get("overtime_hours") else "нет"),("Фактическая стоимость",f"{cost.get('cost_with_overtime',0):,.2f} ₽".replace(","," ")),("Рекомендуемый штат",f"{cost.get('needed_employees',0)} чел."),("Минимально выполнимый",f"{cost.get('staffing_minimum_feasible') or '—'} чел."),("Без переработки",f"{cost.get('staffing_zero_overtime') or '—'} чел."),("Рекомендация",cost.get("recommendation","—"))]
    table=doc.add_table(rows=1,cols=2); table.style="Table Grid"; table.rows[0].cells[0].text="Показатель"; table.rows[0].cells[1].text="Значение"; _style_header(table.rows[0])
    for a,b in rows:
        row=table.add_row(); row.cells[0].text=a; row.cells[1].text=str(b); _set_cant_split(row)
    _protect_table(table); return table


def _add_compact_schedule_table(cell, tasks, project):
    show_floor = len(project.floors) > 1
    headers = (['Этаж'] if show_floor else []) + ['№', 'Комната', 'Тип', 'м²', 'Время', 'мин.']
    table = cell.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    table.autofit = True
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
    _style_header(table.rows[0])
    for t in sorted(tasks, key=lambda x: x.start_dt):
        room = _find_room(project, t.room_id, t.floor_index)
        row = table.add_row()
        _set_cant_split(row)
        offset = 0
        values = []
        if show_floor:
            values.append(str(t.floor_index + 1))
        values.extend([
            str(t.room_id + 1),
            room.name if room else '—',
            room.room_type if room else '—',
            f'{room.area_m2:.1f}' if room else '—',
            f'{t.start_dt:%H:%M}–{t.end_dt:%H:%M}',
            str(int(round((t.end_dt - t.start_dt).total_seconds() / 60))),
        ])
        for i, v in enumerate(values):
            row.cells[i].text = v
            for par in row.cells[i].paragraphs:
                par.paragraph_format.space_after = Pt(0)
                for run in par.runs:
                    run.font.size = Pt(7.5)
        if getattr(t, 'is_overtime', False):
            for c in row.cells:
                _set_cell_shading(c, 'F4CCCC')
    _protect_table(table)
    return table


def _staffing_diagnostics_paragraph(doc, cost):
    tested = cost.get('staffing_tested') or {}
    if not tested:
        return
    p = doc.add_paragraph()
    p.add_run('Проверка штатности: ').bold = True
    minimum = cost.get('staffing_minimum_feasible')
    zero = cost.get('staffing_zero_overtime')
    threshold = cost.get('staffing_overtime_threshold', 30)
    if minimum:
        p.add_run(f'минимально выполнимо: {minimum} чел.; без переработки: {zero or "нет в проверенном диапазоне"}; порог рекомендации: {threshold:.0f} мин. ')
    parts = []
    diagnostics = cost.get('staffing_diagnostics') or {}
    for n, ok in sorted(tested.items()):
        d = diagnostics.get(n, {})
        suffix = ''
        ot = d.get('overtime_minutes', 0)
        miss = d.get('missed_cleanings', 0)
        if ot:
            suffix += f', переработка {int(round(ot))} мин'
        if miss:
            suffix += f', пропущено {miss}'
        parts.append(f'{n} чел. — {"выполнимо" if ok else "нет"}{suffix}')
    p.add_run('; '.join(parts) + '.')


def generate_report(project, filepath, allow_invalid=False):
    validation = validate_schedule(project, schedule_date=getattr(project, 'start_date', None))
    if not validation.get('valid') and not allow_invalid:
        raise ReportGenerationError(
            'Production-экспорт остановлен: расписание невыполнимо.\n\n' +
            (validation.get('violations_summary') or 'Расписание не прошло проверку.')
        )

    cost = calculate_cost(project)
    doc = Document()
    _set_styles(doc)
    _set_margins(doc.sections[0], .4, .4, .45, .45)
    _add_footer(doc.sections[0], project)

    active = [r for r in project.all_rooms() if not getattr(r, 'disabled', False)]
    disabled = [r for r in project.all_rooms() if getattr(r, 'disabled', False)]

    # ---- Титульная / управленческая сводка ----
    doc.add_heading(project.name, level=0)
    doc.add_paragraph('Производственное расписание уборки')
    doc.add_paragraph(f'Дата расписания: {project.start_date:%d.%m.%Y}')
    doc.add_paragraph(
        f'Этажей: {len(project.floors)} • Сотрудников: {project.employees_count} • '
        f'Активных помещений: {len(active)} • Площадь: {sum(r.area_m2 for r in active):.1f} м²'
    )
    sh = project.shifts[0] if project.shifts else None
    lunch = project.breaks[0] if project.breaks else None
    if sh:
        doc.add_paragraph(f'Смена: {sh.start_time}–{sh.end_time}')
    if lunch:
        doc.add_paragraph(f'Обед: {lunch[0]}–{lunch[1]}')

    status = doc.add_paragraph()
    r = status.add_run('● РАСПИСАНИЕ ВЫПОЛНИМО' if validation.get('valid') else '● РАСПИСАНИЕ НЕВЫПОЛНИМО')
    r.bold = True
    r.font.size = Pt(14)
    r.font.color.rgb = RGBColor(24, 122, 59) if validation.get('valid') else RGBColor(176, 0, 32)

    doc.add_paragraph(
        f'Погода: {getattr(project, "weather_factor", 1.0)} • '
        f'Тип уборки: {getattr(project, "cleaning_type", "поддерживающая")}'
    )
    doc.add_heading('Итоги', level=1)
    _add_summary_table(doc, cost)
    _staffing_diagnostics_paragraph(doc, cost)
    if not validation.get('valid'):
        doc.add_paragraph('Диагностический экспорт. Документ НЕ предназначен для передачи в работу.')

    if disabled:
        doc.add_heading('Исключённые помещения', level=2)
        for room in disabled:
            fi = next((i for i, f in enumerate(project.floors) if room in f.rooms), 0)
            doc.add_paragraph(f'Этаж {fi + 1}, №{room.id + 1} {room.name} — исключено пользователем', style='List Bullet')

    # ---- Общие планы ----
    doc.add_page_break()
    doc.add_heading('Зоны ответственности', level=1)
    for fi, floor in enumerate(project.floors):
        if not floor.rooms:
            continue
        if fi:
            doc.add_page_break()
        doc.add_heading(f'{floor.name} • площадь {_floor_area(floor):.1f} м²', level=2)
        stream = _draw_floor_plan(project, floor)
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(stream, **_fit_picture(stream, 9.7, 5.6))
        legend = doc.add_table(rows=1, cols=5)
        legend.style = 'Table Grid'
        for i, h in enumerate(('Сотрудник', 'Комнат', 'Площадь', 'Комнаты', 'Цвет')):
            legend.rows[0].cells[i].text = h
        _style_header(legend.rows[0])
        for emp in range(project.employees_count):
            rooms = [r for r in floor.rooms if (z := _zone_for(project, fi, r.id)) and z.employee_index == emp]
            if not rooms:
                continue
            z = _zone_for(project, fi, rooms[0].id)
            row = legend.add_row(); _set_cant_split(row)
            vals = [
                project.employee_names[emp] if emp < len(project.employee_names) else f'Сотрудник {emp + 1}',
                str(len(rooms)),
                f'{sum(r.area_m2 for r in rooms):.1f} м²',
                ', '.join(f'№{r.id + 1}' for r in sorted(rooms, key=lambda x: x.id)),
                _hex_color(z.color[:3]),
            ]
            for i, v in enumerate(vals): row.cells[i].text = v
            _set_cell_shading(row.cells[0], _hex_color(z.color[:3]))
        _protect_table(legend)

    # ---- Индивидуальные листы: план слева, расписание справа ----
    for emp in range(project.employees_count):
        name = project.employee_names[emp] if emp < len(project.employee_names) else f'Сотрудник {emp + 1}'
        for fi, floor in enumerate(project.floors):
            zone_rooms = [r for r in floor.rooms if (z := _zone_for(project, fi, r.id)) and z.employee_index == emp]
            if not zone_rooms:
                continue
            tasks = [t for t in project.cleaning_tasks if t.employee == emp and t.floor_index == fi]
            tasks.sort(key=lambda t: t.start_dt)

            # Разбиваем длинную таблицу на конечные страницы, сохраняя план слева.
            chunk_size = 18
            chunks = [tasks[i:i + chunk_size] for i in range(0, len(tasks), chunk_size)] or [[]]
            for chunk_idx, chunk in enumerate(chunks):
                _new_landscape(doc, project)
                doc.add_heading(f'{name} — зона ответственности, {floor.name}', level=1)
                p = doc.add_paragraph()
                p.add_run(
                    f'Дата: {project.start_date:%d.%m.%Y} • Погода: {getattr(project, "weather_factor", 1.0)} • '
                    f'Площадь зоны: {sum(r.area_m2 for r in zone_rooms):.1f} м²'
                )

                outer = doc.add_table(rows=1, cols=2)
                outer.autofit = False
                left, right = outer.rows[0].cells
                left.width = Inches(5.0); right.width = Inches(5.0)
                left.vertical_alignment = right.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP

                left.paragraphs[0].text = 'ПЛАН ЗОНЫ'
                left.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                left.paragraphs[0].runs[0].bold = True
                plan_stream = _draw_floor_plan(project, floor, employee_index=emp)
                left.add_paragraph().add_run().add_picture(plan_stream, **_fit_picture(plan_stream, 4.65, 5.5))

                right.paragraphs[0].text = 'ТАБЛИЦА РАСПИСАНИЯ'
                right.paragraphs[0].runs[0].bold = True
                right.add_paragraph(f'Задач на странице: {len(chunk)}')
                _add_compact_schedule_table(right, chunk, project)

                _protect_table(outer)
                if chunk_idx < len(chunks) - 1:
                    doc.add_paragraph('Продолжение расписания → следующая страница').alignment = WD_ALIGN_PARAGRAPH.RIGHT

    doc.save(filepath)
